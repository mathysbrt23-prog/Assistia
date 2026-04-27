from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "assistia-investor-pitch.docx"

GREEN = "25D366"
BLACK = "050505"
DARK = "111111"
SOFT = "F5F7F6"
LINE = "D9E2DD"
TEXT = RGBColor(28, 31, 29)
MUTED = RGBColor(92, 100, 96)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="FFFFFF", size="0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, pct=100):
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:type"), "pct")
    width.set(qn("w:w"), str(pct * 50))


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = add_run(p, text.upper(), bold=True, color=GREEN, size=8.5)
    run.font.name = "Arial"
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.add_run(text)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Bullet")
        p.add_run(item)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_border(cell, LINE, "8")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=BLACK, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = add_run(p2, text, color="4D5551", size=10)
    run.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_grid(doc):
    data = [
        ("Canal", "WhatsApp Business Cloud API"),
        ("Produit", "Assistant Gmail + Calendar avec confirmations"),
        ("Monetisation", "Abonnements SaaS via Stripe"),
        ("Statut", "MVP fonctionnel, prêt à mesurer l’usage réel"),
    ]
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    for idx, (label, value) in enumerate(data):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        set_cell_shading(cell, SOFT)
        set_cell_border(cell, "FFFFFF", "16")
        set_cell_margins(cell, 180, 180, 180, 180)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        add_run(p, label.upper(), bold=True, color=GREEN, size=8)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, value, bold=True, color=BLACK, size=11)


def add_pricing_table(doc):
    headers = ["Plan", "Prix", "Utilisateur cible", "Valeur principale"]
    rows = [
        ["Découverte", "0 €/mois", "Test individuel", "Valider l’usage avec un volume limité"],
        ["Pro", "9 €/mois", "Freelance / entrepreneur", "Usage quotidien Gmail + Calendar"],
        ["Business", "29 €/mois", "Équipe ou sales", "Volumes plus élevés, historique, support"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLACK)
        set_cell_border(cell, BLACK, "4")
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        add_run(p, header, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_shading(cells[idx], "FFFFFF")
            set_cell_border(cells[idx], LINE, "6")
            set_cell_margins(cells[idx], 140, 120, 140, 120)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            add_run(p, value, bold=idx in (0, 1), color=BLACK if idx in (0, 1) else "4D5551", size=9)


def add_roadmap_table(doc):
    headers = ["Horizon", "Priorités produit"]
    rows = [
        ["Court terme", "Onboarding, plans Stripe alignés, suivi des quotas, analytics produit."],
        ["Produit", "Création d’événements, envoi d’emails confirmé, synthèse quotidienne proactive."],
        ["Expansion", "CRM, mode équipe, multi-comptes Google, connecteurs Slack ou Outlook."],
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 100)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLACK)
        set_cell_border(cell, BLACK, "4")
        set_cell_margins(cell, 130, 150, 130, 150)
        add_run(cell.paragraphs[0], header, bold=True, color="FFFFFF", size=9)
    for horizon, value in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((horizon, value)):
            set_cell_shading(cells[idx], "FFFFFF")
            set_cell_border(cells[idx], LINE, "6")
            set_cell_margins(cells[idx], 140, 150, 140, 150)
            add_run(cells[idx].paragraphs[0], text, bold=idx == 0, color=BLACK if idx == 0 else "4D5551", size=9)


def add_section(doc, number, title, paragraphs, bullets=None):
    add_kicker(doc, f"{number:02d} / {title}")
    add_h1(doc, title)
    for paragraph in paragraphs:
        add_body(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = TEXT

    body = styles.add_style("Body", 1)
    body.base_style = normal
    body.font.name = "Arial"
    body.font.size = Pt(10.2)
    body.font.color.rgb = TEXT
    body.paragraph_format.line_spacing = 1.16
    body.paragraph_format.space_after = Pt(7)

    bullet = styles.add_style("Bullet", 1)
    bullet.base_style = body
    bullet.paragraph_format.left_indent = Cm(0.45)
    bullet.paragraph_format.first_line_indent = Cm(-0.18)
    bullet.paragraph_format.space_after = Pt(4)

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(19)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLACK)

    h2 = styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLACK)


def build():
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "Assistia - Pitch investisseur"
    doc.core_properties.subject = "Document de présentation investisseur"
    doc.core_properties.author = "Assistia"

    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    cover = doc.add_table(rows=1, cols=1)
    set_table_width(cover, 100)
    cover_cell = cover.cell(0, 0)
    set_cell_shading(cover_cell, BLACK)
    set_cell_border(cover_cell, BLACK, "0")
    set_cell_margins(cover_cell, 560, 520, 560, 520)
    p = cover_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(34)
    add_run(p, "ASSISTIA", bold=True, color=GREEN, size=12)
    title = cover_cell.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    add_run(title, "Pitch investisseur", bold=True, color="FFFFFF", size=31)
    subtitle = cover_cell.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(26)
    add_run(
        subtitle,
        "L’assistant personnel conversationnel dans WhatsApp pour gérer emails, calendrier et actions courantes.",
        color="D6DED9",
        size=14,
    )
    note = cover_cell.add_paragraph()
    add_run(note, "MVP Next.js · Supabase · Stripe · Google · WhatsApp · OpenAI", color="A7B0AC", size=10)

    doc.add_paragraph()
    add_metric_grid(doc)
    doc.add_page_break()

    add_section(
        doc,
        1,
        "Vision",
        [
            "Assistia veut devenir l’assistant personnel conversationnel des professionnels qui vivent déjà dans WhatsApp.",
            "Le produit place une couche d’assistance au-dessus des outils existants : l’utilisateur écrit une demande, Assistia lit les informations autorisées, résume, prépare une action et demande confirmation quand l’action est sensible.",
        ],
        [
            "Les usages professionnels sur WhatsApp sont déjà installés.",
            "Les APIs Gmail, Calendar, WhatsApp et Stripe permettent un MVP complet.",
            "Les modèles IA rendent possible une interface naturelle sans commandes complexes.",
        ],
    )

    add_section(
        doc,
        2,
        "Problème",
        [
            "Une journée de travail simple est fragmentée entre Gmail, Calendar, Slack, Outlook, WhatsApp et plusieurs dashboards.",
            "Chaque micro-tâche impose de changer d’interface : retrouver un email, vérifier un horaire, déplacer un rendez-vous ou préparer une réponse client.",
        ],
        [
            "Perte de temps sur des tâches de coordination.",
            "Retard dans les réponses client.",
            "Charge mentale élevée pour les indépendants et petites équipes.",
            "Trop d’outils ajoutent une app de plus au lieu de réduire la friction.",
        ],
    )

    add_section(
        doc,
        3,
        "Solution",
        [
            "Assistia est un assistant personnel accessible depuis WhatsApp. L’utilisateur pose une question ou demande une action, puis Assistia répond dans le même canal.",
            "WhatsApp est l’avantage clé : pas de nouvelle application à apprendre, une expérience mobile naturelle et une interaction rapide par conversation.",
        ],
        [
            "Consulter l’agenda du jour ou de l’après-midi.",
            "Résumer les emails du jour ou les emails importants.",
            "Préparer une réponse à un email.",
            "Proposer un déplacement de réunion avec confirmation explicite.",
        ],
    )

    add_section(
        doc,
        4,
        "Produit",
        [
            "Le MVP couvre la chaîne complète : landing, inscription, connexion des outils, abonnement, webhook WhatsApp, agent IA, confirmations et logs.",
            "Les fonctions Gmail incluent la récupération des emails du jour, des emails non lus, des emails importants, la recherche et la préparation de réponses.",
            "Les fonctions Calendar incluent les événements du jour, les événements de l’après-midi, la recherche d’événement et le déplacement confirmé.",
        ],
    )
    add_callout(
        doc,
        "Exemple concret",
        "“Décale mon rendez-vous avec Paul à demain 15h” : Assistia trouve l’événement, prépare la modification et demande “Confirmer ? Réponds OUI pour valider.”",
    )

    add_section(
        doc,
        5,
        "Démo utilisateur",
        [
            "Le matin, l’utilisateur écrit : “Envoie-moi mon planning de la journée”. Assistia répond avec les réunions prévues et les emails importants.",
            "Plus tard, l’utilisateur demande : “Tu as reçu le mail de Léo ?”. Assistia recherche dans Gmail et résume l’information utile.",
            "Si l’utilisateur demande une réponse, Assistia génère un brouillon. Si l’utilisateur demande une modification Calendar, Assistia attend une confirmation explicite.",
        ],
    )

    add_section(
        doc,
        6,
        "Marché",
        [
            "La cible initiale regroupe les freelances, entrepreneurs, dirigeants de petites structures, commerciaux, consultants et profils mobiles qui utilisent déjà WhatsApp comme canal de travail.",
            "Le marché de départ est direct : professionnels qui veulent gagner du temps sans adopter une suite enterprise lourde.",
        ],
        [
            "Résumé de la journée.",
            "Priorisation des emails.",
            "Préparation de réponses client.",
            "Vérification et déplacement de rendez-vous.",
            "Assistant léger pour utilisateurs sans assistant humain.",
        ],
    )

    add_section(
        doc,
        7,
        "Business model",
        [
            "Assistia est monétisable par abonnement SaaS. Le projet intègre déjà Stripe Checkout, Stripe Billing Portal, webhooks de synchronisation, blocage de l’agent si l’abonnement n’est pas actif et quotas mensuels par plan.",
            "La logique commerciale peut commencer avec une offre gratuite pour tester, puis deux plans payants selon le volume et les besoins avancés.",
        ],
    )
    add_pricing_table(doc)

    add_section(
        doc,
        8,
        "Différenciation",
        [
            "Assistia ne demande pas d’utiliser une app de plus : l’usage quotidien se fait dans WhatsApp.",
            "Le dashboard sert à configurer, connecter et auditer, mais la valeur est délivrée dans la conversation.",
            "Le produit est conversationnel, centralisé et contrôlé : l’utilisateur exprime une intention naturelle, Assistia choisit l’outil adapté et les actions sensibles demandent validation.",
        ],
    )

    add_section(
        doc,
        9,
        "Traction",
        [
            "Assistia dispose d’un MVP fonctionnel avec landing, auth, dashboard, Supabase, Stripe, Google OAuth, Gmail, Calendar, WhatsApp webhook, agent IA, confirmations et RLS.",
            "Si aucune donnée utilisateur n’est encore disponible, les premières métriques à suivre sont l’activation, les connexions Google, les numéros WhatsApp associés, la fréquence des demandes et la conversion payante.",
        ],
    )
    add_callout(
        doc,
        "Signal à prouver",
        "Le signal fort n’est pas seulement l’inscription : c’est la fréquence d’usage, c’est-à-dire plusieurs demandes WhatsApp par semaine et par utilisateur actif.",
    )

    add_section(
        doc,
        10,
        "Roadmap",
        [
            "Les prochaines étapes doivent renforcer l’activation, la mesure d’usage et la profondeur produit tout en gardant le principe de confirmation pour les actions sensibles.",
        ],
    )
    add_roadmap_table(doc)

    add_section(
        doc,
        11,
        "Tech stack",
        [
            "Next.js sert l’application web, les pages et les routes API. Supabase gère l’authentification, la base de données, les profils, messages, confirmations et logs avec RLS.",
            "Stripe gère les abonnements. Google OAuth connecte Gmail et Calendar. WhatsApp Business Cloud API reçoit et envoie les messages. OpenAI comprend les demandes et génère les résumés ou brouillons.",
        ],
    )

    add_section(
        doc,
        12,
        "Vision long terme",
        [
            "Assistia commence par emails et calendrier depuis WhatsApp, mais la vision long terme est plus large : devenir un OS personnel conversationnel.",
            "L’utilisateur ne devrait pas ouvrir cinq applications pour savoir quoi faire, retrouver une information ou préparer une action. Assistia peut devenir la couche d’exécution personnelle au-dessus des outils professionnels existants.",
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(footer, "Assistia · Pitch investisseur", color="7A827E", size=8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
