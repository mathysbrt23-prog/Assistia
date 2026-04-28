from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "product" / "assistia-reply-analyse-produit.docx"
LOGO = ROOT / "public" / "assistia-logo.png"

BLACK = "0B0D0F"
CHARCOAL = "15181C"
GREEN = "25D366"
GREEN_DARK = "0F6F3D"
WHITE = "FFFFFF"
TEXT = "151515"
LIGHT_BG = "F4F6F5"
BORDER = "D9E2DC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "8") -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color, size)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Assistia Reply - Analyse produit | ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(120, 120, 120)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    add_page_number_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor(30, 30, 30)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 28, WHITE, 0, 10),
        ("Heading 1", 18, TEXT, 18, 7),
        ("Heading 2", 13, TEXT, 12, 5),
        ("Heading 3", 11, GREEN_DARK, 9, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = name != "Heading 3"
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, BLACK, "0")
    set_cell_margins(cell, top=520, start=520, bottom=520, end=520)

    if LOGO.exists():
        p_logo = cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.add_run().add_picture(str(LOGO), width=Cm(2.15))
    else:
        cell.paragraphs[0].add_run("Assistia").font.color.rgb = RGBColor.from_string(WHITE)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Assistia Reply")
    r.font.name = "Arial"
    r.font.size = Pt(33)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Analyse produit & plan de mise en place")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(210, 214, 212)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Transformer Assistia en assistant de réponses intégré directement dans les mails et conversations, "
        "sans obliger l’utilisateur à changer d’application."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(190, 196, 193)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Document de décision - Avril 2026")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(155, 163, 158)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BG, accent: str = GREEN) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(0.25)
    table.columns[1].width = Cm(16.0)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_shading(left, accent)
    set_cell_shading(right, fill)
    set_cell_border(left, accent, "0")
    set_cell_border(right, fill, "0")
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 165, 200, 165, 200)
    p = right.paragraphs[0]
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(TEXT)
    p2 = right.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def set_col_widths(row, widths_cm: list[float]) -> None:
    for cell, width in zip(row.cells, widths_cm):
        cell.width = Cm(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    if widths_cm:
        set_col_widths(hdr, widths_cm)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, CHARCOAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        if widths_cm:
            set_col_widths(table.rows[-1], widths_cm)
        for i, text in enumerate(row):
            c = cells[i]
            if i == 0:
                set_cell_shading(c, "F8FAF9")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(text) > 16 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(9)
            if text.startswith("Recommandé") or text in ("Faible", "Fort", "Très fort"):
                r.font.bold = True
    set_table_borders(table)
    doc.add_paragraph()


def add_source(doc: Document, title: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title + " - ")
    r.font.bold = True
    r.font.size = Pt(8.8)
    u = p.add_run(url)
    u.font.size = Pt(8.8)
    u.font.color.rgb = RGBColor.from_string(GREEN_DARK)


def build_doc() -> None:
    doc = Document()
    set_document_defaults(doc)
    add_cover(doc)

    doc.add_heading("1. Verdict", level=1)
    add_callout(
        doc,
        "Recommandation principale",
        "Oui, l’idée est plus réaliste que l’agent WhatsApp complet. Le meilleur angle est de repositionner Assistia comme un assistant de réponses intégré : il génère une réponse à partir du contexte, ou reformule ce que l’utilisateur veut dire dans un style professionnel, directement dans Gmail, Outlook, LinkedIn, WhatsApp Web et, plus tard, sur mobile via clavier.",
    )
    doc.add_paragraph(
        "Le produit répond à un besoin court, fréquent et facile à comprendre : écrire mieux et plus vite sans quitter la conversation. "
        "Contrairement à un agent WhatsApp qui doit lire Gmail, modifier le calendrier, gérer OAuth, confirmer des actions et maintenir une logique multi-outils complexe, Assistia Reply commence par une action simple : produire du texte que l’utilisateur valide lui-même."
    )

    doc.add_heading("2. Ce que devient Assistia", level=1)
    doc.add_paragraph(
        "Assistia devient un copilote de réponses. Il vit à côté du champ de texte, pas dans une application séparée. L’utilisateur peut soit demander une réponse complète, soit écrire une intention brute et laisser Assistia la reformuler."
    )
    add_bullets(
        doc,
        [
            "Mode réponse : Assistia lit le message visible ou le texte sélectionné, puis propose une réponse prête à coller.",
            "Mode reformulation : l’utilisateur écrit une idée simple, par exemple “dis-lui que je suis intéressé mais que le prix est trop haut”, et Assistia transforme cela en réponse claire, polie et professionnelle.",
            "Modes de ton : professionnel, direct, chaleureux, ferme, court, commercial.",
            "Validation humaine obligatoire : Assistia insère un brouillon, mais n’envoie jamais le message à la place de l’utilisateur.",
            "Positionnement : “écris la bonne réponse, au bon endroit, sans changer d’app”.",
        ],
    )

    doc.add_heading("3. Pourquoi cette idée est plus forte", level=1)
    add_table(
        doc,
        ["Critère", "Agent WhatsApp initial", "Assistia Reply"],
        [
            ["Complexité", "Très élevée : OAuth, Gmail, Calendar, WhatsApp API, confirmations, webhooks.", "Moyenne : extension navigateur + backend IA + abonnement."],
            ["Usage", "L’utilisateur doit penser à écrire à un bot.", "L’utilisateur voit Assistia exactement là où il répond déjà."],
            ["Risque concurrentiel", "Fort : les hébergeurs et suites IA peuvent proposer un agent assistant généraliste.", "Moins frontal : expérience ultra simple, intégrée aux conversations."],
            ["Valeur perçue", "Assistant global, mais promesse large et difficile à tenir.", "Gain immédiat sur une douleur précise : répondre vite, bien, sans stress."],
            ["MVP", "Long et fragile.", "Possible rapidement sur Gmail et WhatsApp Web."],
        ],
        [3.2, 6.7, 6.7],
    )

    doc.add_heading("4. MVP recommandé", level=1)
    doc.add_paragraph(
        "Le MVP ne doit pas commencer par une application mobile complète. Il doit commencer sur ordinateur, sous forme d’extension Chrome Manifest V3, parce que c’est le chemin le plus direct pour s’intégrer dans Gmail, Outlook Web, LinkedIn et WhatsApp Web."
    )
    add_numbered(
        doc,
        [
            "Extension Chrome : bouton Assistia dans les zones de réponse Gmail et WhatsApp Web.",
            "Deux actions visibles : “Générer une réponse” et “Reformuler mon brouillon”.",
            "Panneau léger : choix du ton, longueur, langue, contexte optionnel.",
            "Insertion du texte généré dans le champ de réponse après validation utilisateur.",
            "Backend Next.js : appel OpenAI, contrôle de l’abonnement Stripe, limites d’usage, logs minimaux.",
            "Dashboard existant : abonnement, usage mensuel, préférences de ton, confidentialité.",
        ],
    )

    doc.add_heading("5. Expérience utilisateur cible", level=1)
    doc.add_paragraph("Exemple concret sur Gmail :")
    add_bullets(
        doc,
        [
            "L’utilisateur ouvre un email client.",
            "Assistia apparaît discrètement près du bouton Répondre.",
            "Il clique sur “Générer une réponse”.",
            "Assistia propose une réponse courte et professionnelle.",
            "L’utilisateur ajuste ou accepte.",
            "Le texte est inséré comme brouillon dans Gmail. L’utilisateur garde le contrôle de l’envoi.",
        ],
    )
    doc.add_paragraph("Exemple concret sur WhatsApp Web ou LinkedIn :")
    add_bullets(
        doc,
        [
            "L’utilisateur sélectionne le dernier message reçu ou clique dans le champ de réponse.",
            "Il écrit son intention en langage naturel : “réponds que je peux demain à 15h, ton sympa”.",
            "Assistia reformule, puis insère la réponse dans la conversation.",
            "L’utilisateur relit et envoie lui-même.",
        ],
    )

    doc.add_heading("6. Faisabilité par plateforme", level=1)
    add_table(
        doc,
        ["Plateforme", "Faisabilité", "Ce qui est simple", "Ce qui est difficile"],
        [
            ["Gmail web", "Forte", "Injecter un bouton via extension, lire une partie visible du fil, insérer un brouillon.", "Gérer les changements fréquents du DOM Gmail et éviter les bugs d’insertion."],
            ["Outlook web", "Forte", "Même logique que Gmail avec content script.", "Sélecteurs et comportements différents selon comptes Microsoft."],
            ["WhatsApp Web", "Bonne", "Lire les messages visibles et insérer une réponse dans le champ.", "WhatsApp Web peut changer son interface, et il faut éviter toute automatisation d’envoi."],
            ["LinkedIn", "Bonne", "Aider à répondre aux messages et prospects.", "Risque de règles anti-automation si le produit ressemble à un bot d’envoi massif."],
            ["iPhone natif", "Moyenne à difficile", "Clavier iOS Assistia ou extension de partage.", "Activation du clavier dans Réglages, accès réseau avec “Full Access”, limitations iOS, confiance utilisateur."],
            ["Android natif", "Moyenne", "Clavier Android avec suggestions et insertion de texte.", "Onboarding clavier, privacy, QA sur beaucoup d’appareils."],
        ],
        [2.6, 2.4, 5.2, 6.2],
    )

    doc.add_heading("7. Choix technique recommandé", level=1)
    doc.add_heading("Version 1 : extension navigateur", level=2)
    doc.add_paragraph(
        "Chrome Manifest V3 permet d’injecter des content scripts sur des sites ciblés comme Gmail ou WhatsApp Web. Le script peut détecter les zones de texte, afficher un petit bouton Assistia et communiquer avec le service worker de l’extension, puis avec le backend Assistia."
    )
    add_bullets(
        doc,
        [
            "Content scripts limités aux domaines nécessaires : mail.google.com, outlook.live.com, web.whatsapp.com, linkedin.com.",
            "Service worker extension : gestion session, appels API, stockage local sécurisé de paramètres non sensibles.",
            "API Next.js : /api/reply/generate, /api/reply/rewrite, /api/usage.",
            "Supabase : profil utilisateur, abonnement, préférences de ton, événements d’usage.",
            "Stripe : Free, Pro, Team avec limites mensuelles.",
            "OpenAI : génération et reformulation avec prompts maîtrisés.",
        ],
    )

    doc.add_heading("Version 2 : mobile", level=2)
    doc.add_paragraph(
        "Le mobile doit arriver après validation du MVP desktop. Sur iPhone, la meilleure option vraiment intégrée est un clavier personnalisé, mais c’est plus sensible : Apple impose des limites aux claviers tiers et l’utilisateur doit explicitement autoriser le clavier. Sur Android, un clavier via InputMethodService est plus flexible, mais la confiance et la confidentialité restent le vrai sujet."
    )

    doc.add_heading("8. Ce qui est simple, moyen, difficile", level=1)
    add_table(
        doc,
        ["Niveau", "Éléments", "Pourquoi"],
        [
            ["Simple", "Repositionner la landing, garder Supabase/Stripe, créer une API de génération, créer des templates de ton.", "Le projet actuel contient déjà le socle SaaS et l’identité visuelle."],
            ["Simple", "Créer un prototype Chrome pour Gmail avec bouton + insertion de brouillon.", "Pas besoin de Gmail API au départ si l’on agit sur le texte visible et le champ de réponse."],
            ["Moyen", "Authentifier une extension avec le compte Assistia.", "Il faut relier proprement session web, extension et limites d’abonnement."],
            ["Moyen", "Supporter plusieurs apps web.", "Chaque app a sa structure DOM, ses changements, ses cas limites."],
            ["Difficile", "iPhone intégré dans toutes les conversations.", "Le clavier iOS est possible mais limité, nécessite Full Access pour appeler l’IA et crée une friction de confiance."],
            ["Difficile", "Extraction fiable du contexte dans toutes les apps mobiles.", "Les apps natives ne donnent pas toutes accès au contenu autour du champ de texte."],
            ["Très difficile", "Envoi automatique ou action autonome.", "À éviter : risque produit, confiance, conformité et erreurs coûteuses."],
        ],
        [2.5, 5.7, 8.2],
    )

    doc.add_heading("9. Architecture produit", level=1)
    doc.add_paragraph("Architecture cible en V1 :")
    add_bullets(
        doc,
        [
            "Frontend existant : landing, pricing, auth, dashboard, checkout Stripe.",
            "Extension Chrome : interface Assistia dans les conversations web.",
            "Backend Next.js : endpoints IA et usage.",
            "Supabase : users, subscriptions, reply_preferences, usage_events, optional_saved_templates.",
            "OpenAI : génération de réponses, reformulation, adaptation du ton.",
            "Aucune action sensible par défaut : pas d’envoi automatique, pas d’accès complet Gmail, pas de modification de calendrier.",
        ],
    )
    add_callout(
        doc,
        "Décision importante",
        "Pour le MVP, éviter l’OAuth Gmail et l’API Gmail. L’extension peut fonctionner avec le texte visible ou sélectionné et insérer un brouillon. Cela réduit fortement la complexité, les scopes sensibles et la friction de confiance.",
        fill="ECFFF3",
        accent=GREEN,
    )

    doc.add_heading("10. Positionnement et pricing", level=1)
    doc.add_paragraph(
        "Le pricing doit rester simple, car le produit vend un gain de temps quotidien et une meilleure qualité de réponse."
    )
    add_table(
        doc,
        ["Plan", "Prix", "Cible", "Contenu"],
        [
            ["Free", "0 EUR/mois", "Découverte", "20 réponses/mois, Gmail + WhatsApp Web, ton professionnel."],
            ["Pro", "9 EUR/mois", "Freelances, sales, entrepreneurs", "Réponses illimitées raisonnables, tons avancés, historique local, raccourcis."],
            ["Team", "29 EUR/mois", "Petites équipes", "Templates partagés, ton de marque, gestion équipe, priorités support."],
        ],
        [2.6, 2.4, 4.1, 7.4],
    )

    doc.add_heading("11. Différenciation", level=1)
    add_bullets(
        doc,
        [
            "Pas une app de plus : Assistia apparaît dans l’outil déjà utilisé.",
            "Pas un agent autonome : l’utilisateur contrôle chaque réponse.",
            "Pas un assistant généraliste vague : un cas d’usage très précis, très fréquent, très monétisable.",
            "Accent français/professionnel : reformulation naturelle pour mails, prospection, relation client, relance.",
            "Confidentialité claire : pas de stockage du contenu des messages par défaut.",
        ],
    )

    doc.add_heading("12. Risques", level=1)
    add_table(
        doc,
        ["Risque", "Impact", "Réponse"],
        [
            ["Concurrence de Grammarly, Gmail/Google, Outlook/Microsoft", "Fort", "Se concentrer sur simplicité, français, multi-apps, ton pro, workflows commerciaux."],
            ["Changements DOM des apps web", "Moyen", "Commencer avec Gmail + WhatsApp Web, tests automatisés, sélecteurs robustes, fallback sélection de texte."],
            ["Confiance sur le contenu privé", "Fort", "Aucune conservation par défaut, consentement explicite, logs sans contenu, politique claire."],
            ["Friction mobile clavier", "Fort", "Ne pas commencer par mobile natif. Tester mobile avec partage/copie, puis clavier si la demande est prouvée."],
            ["Extension store review", "Moyen", "Permissions minimales, description claire, pas de code distant exécuté dans l’extension."],
        ],
        [4.5, 2.5, 9.5],
    )

    doc.add_heading("13. Roadmap réaliste", level=1)
    add_table(
        doc,
        ["Période", "Objectif", "Livrable"],
        [
            ["Semaine 1", "Cadrage produit", "Landing repositionnée, prompts, maquettes extension, pricing."],
            ["Semaine 2", "Prototype Gmail", "Bouton Assistia, génération, reformulation, insertion brouillon."],
            ["Semaine 3", "SaaS", "Auth extension, Stripe, limites d’usage, dashboard usage."],
            ["Semaine 4", "Beta", "WhatsApp Web, onboarding, feedback, premiers utilisateurs."],
            ["Mois 2", "Distribution", "Chrome Web Store, Outlook/LinkedIn, templates métiers."],
            ["Mois 3", "Mobile test", "Prototype clavier Android ou iOS, selon traction."],
        ],
        [2.4, 4.2, 9.8],
    )

    doc.add_heading("14. Métriques de validation", level=1)
    add_bullets(
        doc,
        [
            "Activation : pourcentage d’utilisateurs qui génèrent au moins 3 réponses dans la première journée.",
            "Rétention : utilisateurs actifs après 7 jours.",
            "Usage : réponses générées par utilisateur et par semaine.",
            "Conversion : passage Free vers Pro.",
            "Qualité : taux de réponses insérées sans modification majeure.",
            "Douleur réelle : combien d’utilisateurs déclarent que l’outil leur évite de repousser des réponses importantes.",
        ],
    )

    doc.add_heading("15. Décision finale", level=1)
    add_callout(
        doc,
        "Plan conseillé",
        "Garder le front Assistia, abandonner temporairement la promesse d’agent WhatsApp complet, et lancer Assistia Reply comme extension desktop. Le produit est plus simple, plus crédible, plus rapide à tester et beaucoup plus proche d’un achat impulsif à 9 EUR/mois.",
        fill="ECFFF3",
        accent=GREEN,
    )
    doc.add_paragraph(
        "La version mobile doit rester dans la vision, mais pas dans le MVP. Le bon ordre est : 1) prouver que les utilisateurs paient pour mieux répondre sur ordinateur ; 2) étendre à plusieurs apps web ; 3) construire le clavier mobile uniquement si l’usage quotidien est confirmé."
    )

    doc.add_heading("Sources techniques utilisées", level=1)
    add_source(doc, "Chrome Extensions - Content scripts", "https://developer.chrome.com/docs/extensions/develop/concepts/content-scripts")
    add_source(doc, "Chrome Extensions - Manifest V3", "https://developer.chrome.com/docs/extensions/develop/migrate/what-is-mv3")
    add_source(doc, "Google Workspace Add-ons - Gmail compose UI", "https://developers.google.com/workspace/add-ons/gmail/extending-compose-ui")
    add_source(doc, "Apple - Custom Keyboard Extension", "https://developer.apple.com/library/archive/documentation/General/Conceptual/ExtensibilityPG/CustomKeyboard.html")
    add_source(doc, "Apple - RequestsOpenAccess for custom keyboards", "https://developer.apple.com/library/archive/documentation/General/Reference/InfoPlistKeyReference/Articles/AppExtensionKeys.html")
    add_source(doc, "Android - InputMethodService", "https://developer.android.com/reference/android/inputmethodservice/InputMethodService.html")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
