from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "investor" / "assistia-reply-pitch-investisseur.docx"
LOGO = ROOT / "public" / "assistia-logo.png"

BLACK = "0B0D0F"
CHARCOAL = "171A1E"
GREEN = "25D366"
GREEN_DARK = "0B7A3F"
LIGHT = "F4F7F5"
BORDER = "D9E2DC"
TEXT = "161A17"
MUTED = "5D6660"
WHITE = "FFFFFF"


SOURCES = [
    ("Radicati Group", "Email Statistics Report 2024-2028 Executive Summary", "https://www.radicati.com/wp/wp-content/uploads/2024/10/Email-Statistics-Report-2024-2028-Executive-Summary.pdf"),
    ("Eurostat", "20% of EU enterprises use AI technologies, 2025", "https://ec.europa.eu/eurostat/web/products-eurostat-news/w/ddn-20251211-2"),
    ("OECD / INSEE", "Financing SMEs and Entrepreneurs 2024 - France", "https://www.oecd.org/en/publications/financing-smes-and-entrepreneurs-2024_fa521246-en/full-report/component-24.html"),
    ("DataReportal", "Digital 2026: France", "https://datareportal.com/reports/digital-2026-france"),
    ("Grand View Research", "AI Assistant Software Market, 2025-2033", "https://www.grandviewresearch.com/industry-analysis/ai-assistant-software-market-report"),
    ("Technavio", "AI Writing Assistant Software Market, 2025-2029", "https://www.technavio.com/report/ai-writing-assistant-software-market-industry-analysis"),
    ("Grammarly Support", "Grammarly Pro pricing", "https://support.grammarly.com/hc/en-us/articles/115000090011-How-much-does-Grammarly-Pro-cost"),
    ("Superhuman Help Center", "Superhuman pricing plans", "https://help.superhuman.com/hc/en-us/articles/38456109456147-Pricing-Plans"),
    ("MailMaestro", "Plans & Pricing", "https://www.maestrolabs.com/pricing"),
    ("Microsoft", "Microsoft 365 Copilot plans and pricing", "https://www.microsoft.com/en-us/microsoft-365-copilot/pricing-new"),
    ("Compose AI", "Pricing", "https://composeai.io/pricing"),
    ("Anthropic", "Claude pricing", "https://claude.com/pricing"),
    ("OpenAI", "ChatGPT pricing", "https://chatgpt.com/pricing/"),
    ("TechCrunch", "WhatsApp has more than 3 billion monthly users", "https://techcrunch.com/2025/05/01/whatsapp-now-has-more-than-3-billion-users/"),
]


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_col_widths(row, widths_cm: list[float]) -> None:
    for cell, width in zip(row.cells, widths_cm):
        cell.width = Cm(width)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Arial"


def style_table(table, header_fill: str = CHARCOAL, header_text: str = WHITE) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = rgb(header_text)
                        run.font.size = Pt(8.8)
            else:
                set_cell_shading(cell, "FFFFFF" if i % 2 else "F8FAF8")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.6)
        if i == 0:
            repeat_header(row)


def footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Assistia Reply - Pitch investisseur | ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after, bold in [
        ("Title", 31, WHITE, 0, 8, True),
        ("Heading 1", 18, TEXT, 14, 6, True),
        ("Heading 2", 13, TEXT, 10, 4, True),
        ("Heading 3", 11, GREEN_DARK, 8, 3, True),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(9.8)
        style.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, BLACK, "0")
    set_cell_margins(cell, top=560, start=540, bottom=560, end=540)
    if LOGO.exists():
        cell.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(2.05))
    else:
        r = cell.paragraphs[0].add_run("Assistia")
        r.font.color.rgb = rgb(WHITE)
        r.font.bold = True

    p = cell.add_paragraph()
    r = p.add_run("Assistia Reply")
    r.font.name = "Arial"
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = rgb(WHITE)

    p = cell.add_paragraph()
    r = p.add_run("Pitch investisseur")
    r.font.name = "Arial"
    r.font.size = Pt(17)
    r.font.color.rgb = rgb("DDE4E0")

    p = cell.add_paragraph()
    r = p.add_run("La bonne réponse, au bon endroit.")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = rgb(GREEN)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(
        "Assistant IA intégré à Gmail et WhatsApp Web pour générer ou reformuler des réponses professionnelles sans changer d’application."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb("B8C3BE")

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Version de travail - Avril 2026")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = rgb("8E9994")
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(0.28)
    table.columns[1].width = Cm(16.4)
    left, right = table.rows[0].cells
    set_cell_shading(left, GREEN)
    set_cell_shading(right, fill)
    set_cell_border(left, GREEN, "0")
    set_cell_border(right, fill, "0")
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 150, 190, 150, 190)
    p = right.paragraphs[0]
    r = p.add_run(title)
    r.font.name = "Arial"
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(TEXT)
    p2 = right.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9.8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
    style_table(table)
    if widths:
        for row in table.rows:
            set_col_widths(row, widths)


def add_toc(doc: Document) -> None:
    doc.add_heading("Sommaire", level=1)
    sections = [
        "Executive summary",
        "Le problème",
        "La solution",
        "Produit - état actuel",
        "Marché",
        "Concurrence",
        "Business model",
        "Go-to-market",
        "Traction & validation",
        "Équipe",
        "Financial projections - 3 ans",
        "Ask & use of funds",
        "Annexes & questions fréquentes investisseurs",
    ]
    for section in sections:
        p = doc.add_paragraph(style="List Number")
        p.add_run(section)
    doc.add_page_break()


def build() -> None:
    doc = Document()
    setup(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("1. EXECUTIVE SUMMARY", level=1)
    add_callout(
        doc,
        "Projet en 3 phrases",
        "Assistia Reply réduit la friction quotidienne des professionnels qui doivent répondre vite, bien et avec le bon ton à leurs clients. Le produit prend la forme d’une extension Chrome intégrée à Gmail et WhatsApp Web : l’utilisateur sélectionne le contexte, écrit son intention, et Assistia propose un brouillon prêt à insérer, sans jamais envoyer à sa place. L’opportunité est de prendre une niche simple, fréquente et monétisable dans le marché en forte croissance des assistants IA d’écriture et de productivité."
    )
    doc.add_heading("Taglines proposées", level=2)
    add_bullets(doc, [
        "La bonne réponse, au bon endroit.",
        "Réponds mieux, sans quitter tes conversations.",
        "Ton copilote de réponse pour les emails et messages clients.",
    ])
    doc.add_heading("Chiffres clés à retenir", level=2)
    add_simple_table(doc, ["Indicateur", "Chiffre", "Lecture investisseur"], [
        ["Volume email mondial", "392,5 milliards d’emails/jour prévus en 2026", "Le problème d’écriture/réponse est massif et quotidien (Radicati Group, 2024-2028)."],
        ["Utilisateurs email mondiaux", "4,73 milliards en 2026", "L’email reste l’infrastructure universelle du travail et du commerce (Radicati Group)."],
        ["Adoption IA en entreprise UE", "20,0% des entreprises 10+ salariés en 2025", "Le marché est passé du test à l’usage professionnel (Eurostat, 2025)."],
        ["Marché AI assistant software", "8,46 Md$ en 2024 ; 35,72 Md$ projetés en 2033", "Vague de fond sur les assistants IA productivité (Grand View Research)."],
        ["Ask recommandé", "150k€ pre-seed ou 500k€ seed", "150k€ pour valider le canal ; 500k€ pour accélérer produit + acquisition."]
    ], [4.0, 4.3, 8.4])

    doc.add_heading("2. LE PROBLÈME", level=1)
    doc.add_paragraph(
        "La cible n’a pas un problème d’accès à l’IA. Elle a un problème de contexte, de friction et de ton. Les professionnels savent ce qu’ils veulent répondre, mais perdent du temps à transformer une intention brute en message clair, crédible et adapté au client."
    )
    doc.add_heading("Pain point concret", level=2)
    add_bullets(doc, [
        "Les emails, messages WhatsApp Web et messages LinkedIn arrivent dans des interfaces séparées.",
        "Le coût cognitif est élevé : comprendre le contexte, trouver le bon angle, rédiger sans paraître froid ou insistant.",
        "Les réponses sensibles sont souvent repoussées : relance client, objection prix, refus, retard, négociation, SAV.",
        "Les outils IA génériques imposent un aller-retour : copier le message, ouvrir ChatGPT, expliquer le contexte, recopier la réponse.",
        "Les assistants intégrés des grandes suites sont puissants mais souvent trop horizontaux, en anglais-first, ou liés à des licences coûteuses."
    ])
    doc.add_heading("Statistiques marché", level=2)
    add_simple_table(doc, ["Statistique", "Source", "Implication"], [
        ["392,5 milliards d’emails envoyés/reçus par jour prévus en 2026.", "Radicati Group, Email Statistics Report 2024-2028.", "Même une micro-amélioration sur la rédaction/réponse crée une grande surface de valeur."],
        ["4,73 milliards d’utilisateurs email mondiaux prévus en 2026.", "Radicati Group.", "L’email reste un canal universel malgré Slack, Teams et WhatsApp."],
        ["20,0% des entreprises européennes de 10+ salariés utilisent l’IA en 2025, contre 13,5% en 2024.", "Eurostat, 2025.", "Le timing est favorable : l’IA entre dans les workflows métiers."],
        ["L’usage IA le plus courant en entreprise UE est l’analyse du langage écrit : 11,8%.", "Eurostat, 2025.", "Assistia se positionne au cœur de l’usage IA déjà adopté."],
        ["La France compte 4,238M d’entreprises en 2022, dont 99,9% de PME.", "OECD / INSEE, 2024.", "Le marché français est massif, fragmenté, et compatible avec une approche bottom-up."],
        ["63,4M d’internautes en France fin 2025, pénétration 95,2%.", "DataReportal, Digital 2026 France.", "Le produit peut être distribué entièrement en ligne, sans vente terrain au départ."]
    ], [5.4, 4.5, 6.8])
    doc.add_heading("Pourquoi les solutions actuelles ne suffisent pas", level=2)
    add_bullets(doc, [
        "ChatGPT/Claude : excellents modèles, mais l’utilisateur sort de son flux de travail.",
        "Grammarly/Superhuman Go : très horizontal, pas pensé pour les cas de réponse client francophones.",
        "Superhuman/Shortwave : puissants mais demandent de remplacer l’inbox, friction majeure pour les indépendants et TPE.",
        "Copilot/Gemini : intégrés mais liés à des suites, à des licences, et moins différenciés sur WhatsApp Web/LinkedIn.",
        "Extensions WhatsApp : souvent low-cost, peu premium, faibles sur le ton professionnel et la confiance."
    ])

    doc.add_heading("3. LA SOLUTION", level=1)
    doc.add_paragraph(
        "Assistia Reply est une couche IA légère qui s’installe là où l’utilisateur répond déjà. Le produit ne veut pas remplacer Gmail, WhatsApp Web ou LinkedIn : il ajoute un bouton et un panneau Assistia dans le champ de réponse."
    )
    doc.add_heading("Les 5 fonctionnalités clés", level=2)
    add_bullets(doc, [
        "Générer une réponse à partir du message visible ou sélectionné.",
        "Reformuler une intention brute en email/message professionnel.",
        "Choisir un ton : professionnel, court, chaleureux, ferme ou commercial.",
        "Insérer le brouillon dans le champ de réponse sans auto-send.",
        "Dashboard SaaS : usage mensuel, préférences, historique léger, abonnement Stripe."
    ])
    doc.add_heading("Différenciation", level=2)
    add_callout(
        doc,
        "Positionnement unique",
        "Assistia n’est pas un chatbot généraliste ni un nouveau client email. C’est un assistant de réponse francophone, contextuel et non-intrusif, centré sur les situations business où la formulation influence la relation client."
    )
    doc.add_heading("User journey type", level=2)
    add_numbered(doc, [
        "L’utilisateur installe l’extension Chrome et se connecte à son compte Assistia.",
        "Dans Gmail ou WhatsApp Web, il ouvre une conversation client.",
        "Il clique sur Assistia, sélectionne un ton et écrit son intention en langage naturel.",
        "Assistia génère un brouillon contextualisé.",
        "L’utilisateur relit, ajuste si besoin, insère et envoie lui-même."
    ])

    doc.add_heading("4. PRODUIT - ÉTAT ACTUEL", level=1)
    doc.add_heading("Stack technique", level=2)
    add_simple_table(doc, ["Bloc", "Technologie", "État"], [
        ["Frontend SaaS", "Next.js App Router, TypeScript, Tailwind CSS", "Landing, pricing, auth, dashboard."],
        ["Auth & Database", "Supabase Auth + Supabase Database + RLS", "Schéma orienté profils, abonnements, demandes de réponses, usage."],
        ["Paiement", "Stripe Checkout, Billing Portal, Webhooks", "Plans Pro/Business configurables par variables d’environnement."],
        ["IA", "OpenAI API", "Route `/api/reply/generate` avec quotas et stockage léger."],
        ["Distribution", "Chrome Extension Manifest V3", "Prototype avec content script Gmail/WhatsApp Web, panneau, génération, insertion."],
        ["Déploiement", "Vercel compatible", "Projet prêt à déployer après env vars et Supabase SQL."]
    ], [3.8, 5.2, 7.7])
    doc.add_heading("État d’avancement honnête", level=2)
    add_bullets(doc, [
        "MVP technique web fonctionnel : landing, dashboard, auth, Stripe, Supabase, API OpenAI.",
        "Prototype extension disponible, mais à fiabiliser sur Gmail avant beta publique.",
        "Pas encore de traction réelle : le prochain objectif est une beta privée de 30 à 50 utilisateurs.",
        "La logique produit a été simplifiée après pivot depuis un agent WhatsApp complet vers un cas d’usage plus testable."
    ])
    doc.add_heading("Roadmap 2026", level=2)
    add_simple_table(doc, ["Période", "Objectif produit", "Livrables"], [
        ["Q2 2026", "Beta Gmail", "Extension stable Gmail, login, génération/reformulation, insertion brouillon, analytics usage."],
        ["Q3 2026", "Gmail + WhatsApp Web", "Support WhatsApp Web, templates métiers, onboarding payant, dashboard usage avancé."],
        ["Q4 2026", "Monétisation & équipes", "Plans Team, ton de marque partagé, templates d’équipe, conformité RGPD, début marketplace Chrome."],
    ], [2.3, 4.4, 9.9])

    doc.add_heading("5. MARCHÉ", level=1)
    doc.add_heading("TAM / SAM / SOM", level=2)
    add_simple_table(doc, ["Niveau", "Chiffrage", "Méthodologie"], [
        ["TAM large", "8,46 Md$ en 2024 ; 35,72 Md$ projetés en 2033.", "Marché mondial AI assistant software (Grand View Research)."],
        ["TAM direct", "+6,21 Md$ d’opportunité 2025-2029 ; CAGR 36%.", "Marché AI writing assistant software (Technavio)."],
        ["SAM France", "≈ 91,6 M€/an [ESTIMATION].", "4,238M entreprises France x 20% cible communication-intensive x 9€/mois x 12. Source base entreprises : OECD/INSEE."],
        ["SAM francophone", "≈ 120-150 M€/an [ESTIMATION].", "SAM France augmenté de 30-60% pour Belgique, Suisse romande, Luxembourg, Québec et marchés francophones premium."],
        ["SOM 3 ans", "≈ 2,3 M€ ARR [ESTIMATION].", "15 000 clients payants x 13€/mois de revenu moyen pondéré x 12 ; soit <3% du SAM France."]
    ], [3.2, 4.0, 9.4])
    doc.add_heading("Croissance et tendances favorables", level=2)
    add_bullets(doc, [
        "Le marché des AI assistants progresse avec un CAGR estimé à 17,5% sur 2025-2033 (Grand View Research).",
        "Le segment AI writing assistant est encore plus rapide : CAGR 36% sur 2024-2029 selon Technavio.",
        "Les entreprises européennes adoptent l’IA par cas d’usage texte avant les automatisations plus complexes (Eurostat).",
        "Les indépendants et TPE veulent des outils à ROI immédiat, sans migration de suite logicielle.",
        "WhatsApp a dépassé 3 milliards d’utilisateurs mensuels en 2025, ce qui confirme l’importance des conversations comme interface de travail (TechCrunch, déclaration Meta)."
    ])

    doc.add_heading("6. CONCURRENCE", level=1)
    add_simple_table(doc, ["Concurrent", "Prix public", "Cible", "Fonctions clés", "Géographie", "Points faibles"], [
        ["Grammarly / Superhuman Go", "$12/mo annuel ou $30/mo mensuel pour Grammarly Pro", "Pros individuels et équipes", "Correction, réécriture, ton, IA partout", "Global", "Très horizontal ; peu spécialisé réponses client francophones."],
        ["Superhuman Mail", "$30/mo Starter ; $40/mo Business", "High-volume email pros", "Client email premium, AI, labels, CRM", "US/global", "Oblige à changer d’inbox ; prix élevé."],
        ["MailMaestro", "Free limité ; $15/seat/mo mensuel", "Professionnels email", "Reply, compose, improve, summarize, scheduling", "Global", "Email-first ; moins fort sur WhatsApp/LinkedIn."],
        ["Microsoft 365 Copilot", "$22-$32/user/mo en bundle SMB promo ou add-on selon licence", "Entreprises Microsoft", "Outlook, Office, Teams, documents", "Global", "Dépendance M365 ; prix/licence complexes."],
        ["Compose AI", "Free + PAYG / plans usage-based", "Utilisateurs Chrome", "Agents, autocomplete, tâches planifiées", "Global", "Positionnement moins centré réponse pro."],
        ["HyperWrite", "Premium/Ultra, prix à revalider selon page publique", "Power users IA writing", "Typeahead, personas, assistant web", "Global", "Apparence plus généraliste ; moins workflow client."],
        ["ChatGPT / Claude", "$20+/mo selon plan", "Grand public et pros", "Génération texte polyvalente", "Global", "Hors contexte ; copier-coller ; pas intégré au champ de réponse."]
    ], [3.0, 2.8, 3.0, 4.2, 2.4, 4.4])
    add_callout(
        doc,
        "Notre positionnement unique",
        "Assistia prend le terrain entre les IA généralistes et les suites enterprise : un produit simple, francophone, intégré au lieu de réponse, orienté messages clients et conçu pour ne jamais envoyer à la place de l’utilisateur."
    )

    doc.add_heading("7. BUSINESS MODEL", level=1)
    add_simple_table(doc, ["Plan", "Prix", "Cible", "Inclus"], [
        ["Free", "0€/mois", "Découverte", "20 réponses/mois, Gmail ou WhatsApp Web, tons de base."],
        ["Pro", "9€/mois", "Freelances, fondateurs, sales indépendants", "1 000 réponses/mois, Gmail + WhatsApp Web, templates relance/prix/refus/SAV, préférences de ton."],
        ["Business", "29€/mois", "Petites équipes commerciales/support", "3 utilisateurs, 3 000 réponses/mois, ton de marque partagé, historique d’usage sans contenu sensible."]
    ], [2.6, 2.5, 4.0, 7.5])
    doc.add_heading("Justification ROI", level=2)
    add_bullets(doc, [
        "[ESTIMATION] Si Assistia économise 10 minutes/jour à un professionnel facturé ou valorisé 35€/h, le gain mensuel vaut environ 117€ pour 20 jours ouvrés.",
        "Le plan Pro à 9€/mois nécessite donc moins de 16 minutes économisées par mois pour être rationnel économiquement.",
        "Pour une équipe commerciale, une seule relance mieux formulée qui convertit un deal justifie plusieurs mois d’abonnement."
    ])
    doc.add_heading("LTV / CAC théorique", level=2)
    add_simple_table(doc, ["Hypothèse", "Valeur", "Commentaire"], [
        ["ARPA payé pondéré", "12€/mois [ESTIMATION]", "Mix Pro majoritaire + Business minoritaire au lancement."],
        ["Marge brute cible", "82% [ESTIMATION]", "Coûts OpenAI + infra + Stripe ; optimisable par caching et modèles moins chers."],
        ["Churn mensuel cible", "<5%", "Objectif SaaS prosumer ; alerte si >7%."],
        ["LTV théorique", "≈ 246€", "ARPA 12€ x marge 82% / churn 4%."],
        ["CAC cible blended", "45-70€", "Nécessaire pour LTV/CAC >3x."]
    ], [4.0, 3.0, 9.7])
    doc.add_heading("Revenus secondaires possibles", level=2)
    add_bullets(doc, [
        "Templates métiers premium : immobilier, agences, freelances, support client.",
        "Plan Team avec ton de marque et bibliothèque partagée.",
        "Offre white-label pour agences ou réseaux de freelances.",
        "Intégrations CRM simples : HubSpot, Pipedrive, Notion CRM."
    ])

    doc.add_heading("8. GO-TO-MARKET", level=1)
    doc.add_heading("Stratégie Year 1", level=2)
    add_simple_table(doc, ["Canal", "Angle", "CAC cible", "Pourquoi"], [
        ["SEO long-tail", "Répondre à un client, relancer devis, objection prix, refus poli", "10-25€", "Intentions de recherche très concrètes, contenu réutilisable en templates."],
        ["TikTok/LinkedIn démos", "Avant/après de réponses clients", "20-50€", "Le produit se comprend en vidéo courte."],
        ["Communautés freelances", "Malt, LinkedIn, Indie Hackers FR, groupes sales", "15-40€", "Cible ressent le problème tous les jours."],
        ["Chrome Web Store SEO", "AI email reply assistant, Gmail reply, WhatsApp Web AI", "5-20€", "Canal naturel pour une extension."],
        ["Partenariats micro-agences", "Pack réponse client pour équipes de 2-10", "60-120€", "Moins scalable, mais ARPA plus haut."],
        ["Paid search test", "requêtes Gmail AI reply / reformuler email pro", "80-150€", "À utiliser seulement après conversion prouvée."]
    ], [3.1, 4.7, 2.5, 6.6])
    doc.add_heading("Funnel cible Year 1", level=2)
    add_simple_table(doc, ["Étape", "Objectif 12 mois", "Taux cible"], [
        ["Visiteurs landing", "120 000", "-"],
        ["Installations extension / signups", "12 000", "10% visitor -> signup"],
        ["Utilisateurs actifs mensuels", "4 800", "40% signup -> MAU"],
        ["Payants", "600", "5% signup -> paid"],
        ["MRR fin Year 1", "≈ 7 200€", "ARPA 12€/mois"],
    ], [4.3, 4.2, 6.0])

    doc.add_heading("9. TRACTION & VALIDATION", level=1)
    doc.add_paragraph(
        "À ce stade, la traction commerciale reste à construire. Le bon message investisseur est donc : MVP construit, problème resserré, expérimentation prête, métriques de validation définies."
    )
    doc.add_heading("Validation en 90 jours", level=2)
    add_simple_table(doc, ["Hypothèse", "Test", "Seuil de validation"], [
        ["Le problème est fréquent", "30 interviews freelances/sales/support", "70% déclarent répondre à des messages clients tous les jours."],
        ["Le produit est utile", "Beta Gmail avec 50 utilisateurs", "30% utilisent Assistia 3 fois/semaine."],
        ["Le prix est acceptable", "Paywall Pro après quota free", "3-5% conversion free -> paid."],
        ["La rétention existe", "Cohorte beta 4 semaines", "Retention W4 >25% des utilisateurs activés."],
        ["La distribution fonctionne", "SEO + démos sociales + communautés", "CAC organique/blended <70€."]
    ], [4.2, 5.2, 6.1])
    doc.add_heading("KPIs à suivre", level=2)
    add_bullets(doc, [
        "Activation : installation extension + première réponse générée.",
        "Usage : réponses générées par utilisateur actif/semaine.",
        "Conversion : free -> Pro, Pro -> Business.",
        "Rétention : W1, W4, M3.",
        "Qualité : taux d’insertion du brouillon, taux de régénération, feedback positif/négatif.",
        "Business : MRR, ARPA, churn, CAC, payback, LTV/CAC."
    ])

    doc.add_heading("10. ÉQUIPE", level=1)
    add_callout(
        doc,
        "Section à compléter par le fondateur",
        "Cette section doit être personnalisée avant envoi. Un investisseur veut comprendre pourquoi toi, maintenant, sur ce problème."
    )
    doc.add_heading("Structure recommandée", level=2)
    add_bullets(doc, [
        "Founder & CEO : parcours, exposition au problème, capacité à vendre et exécuter.",
        "Capacité produit/tech : démontrer que le MVP a été construit vite et avec une stack moderne.",
        "Advisors possibles : sales B2B, growth SaaS, sécurité/RGPD, Chrome extension.",
        "Recrutements 12 mois : full-stack/extension engineer, growth marketer, customer success part-time."
    ])

    doc.add_heading("11. FINANCIAL PROJECTIONS - 3 ANS", level=1)
    doc.add_heading("Hypothèses réalistes", level=2)
    add_bullets(doc, [
        "Conversion freemium -> paid : 3-5% à 12 mois, 5-7% après amélioration onboarding.",
        "Churn mensuel cible : <5% sur Pro, <3% sur Business.",
        "Marge brute : 80-85% après optimisation modèles et quotas.",
        "ARPA payé : 12€ au départ, puis 13-16€ avec hausse part Business.",
        "Coûts principaux : développement extension, marketing, support, OpenAI/API, Stripe, Supabase/Vercel."
    ])
    doc.add_heading("Scénario réaliste - jalons 36 mois", level=2)
    add_simple_table(doc, ["Mois", "Free users", "Paid users", "MRR", "ARR", "Coûts mensuels", "Burn mensuel"], [
        ["M0", "0", "0", "0€", "0€", "1 000€", "1 000€"],
        ["M6", "2 500", "90", "1 080€", "12 960€", "8 000€", "6 920€"],
        ["M12", "12 000", "600", "7 200€", "86 400€", "18 000€", "10 800€"],
        ["M18", "24 000", "1 500", "19 500€", "234 000€", "34 000€", "14 500€"],
        ["M24", "45 000", "3 200", "44 800€", "537 600€", "52 000€", "7 200€"],
        ["M30", "70 000", "5 500", "77 000€", "924 000€", "72 000€", "-5 000€"],
        ["M36", "95 000", "8 000", "112 000€", "1 344 000€", "92 000€", "-20 000€"],
    ], [2.1, 2.3, 2.3, 2.2, 2.4, 3.1, 3.2])
    doc.add_heading("3 scénarios", level=2)
    add_simple_table(doc, ["Scénario", "M36 paid users", "ARPA", "MRR M36", "ARR M36", "Lecture"], [
        ["Conservative", "2 000", "11€", "22 000€", "264 000€", "Produit viable mais distribution lente ; bootstrappable."],
        ["Realistic", "8 000", "14€", "112 000€", "1 344 000€", "SaaS VC-compatible si churn maîtrisé et CAC <70€."],
        ["Optimistic", "25 000", "14€", "350 000€", "4 200 000€", "Accélération forte via Chrome Web Store, SEO et partenariats équipes."]
    ], [3.0, 3.0, 2.1, 2.5, 2.8, 6.1])
    doc.add_paragraph(
        "Break-even visé : M28-M32 dans le scénario réaliste, sous réserve d’un CAC blended inférieur à 70€, d’une marge brute >80% et d’un churn mensuel <5%."
    )

    doc.add_heading("12. ASK & USE OF FUNDS", level=1)
    add_simple_table(doc, ["Scénario", "Montant", "Instrument", "Pre-money suggérée", "Runway", "Milestones"], [
        ["Pre-seed lean", "150k€", "BSA-AIR / SAFE", "1,2-1,8M€", "12 mois", "Beta Gmail, 2 000 signups, 150 payants, MRR 1,5-2k€, preuve d’usage."],
        ["Seed", "500k€", "BSA-AIR puis equity", "2,5-4M€", "18 mois", "Gmail + WhatsApp Web, 50k signups, 2k payants, MRR 20-30k€."],
        ["Seed+", "1M€", "Equity ou SAFE capé", "5-8M€", "18-24 mois", "Équipe 5-7, expansion francophone, MRR 80-120k€, churn prouvé."]
    ], [2.5, 2.3, 3.2, 3.2, 2.0, 5.8])
    doc.add_heading("Usage des fonds", level=2)
    add_simple_table(doc, ["Poste", "Pre-seed 150k€", "Seed 500k€", "Seed+ 1M€"], [
        ["Produit & engineering", "45%", "38%", "32%"],
        ["Marketing & acquisition", "25%", "32%", "36%"],
        ["Salaires fondateurs / équipe", "20%", "22%", "24%"],
        ["Infra, IA, sécurité, juridique", "10%", "8%", "8%"],
    ], [4.5, 3.3, 3.3, 3.3])

    doc.add_heading("ANNEXES & QUESTIONS FRÉQUENTES INVESTISSEURS", level=1)
    doc.add_heading("10 questions difficiles + réponses préparées", level=2)
    qa = [
        ("Pourquoi Google/Microsoft ne vont pas vous écraser ?", "Ils couvrent la suite complète, pas le cas d’usage étroit “réponse client francophone dans Gmail + WhatsApp Web + LinkedIn”. Assistia gagne sur spécialisation, vitesse produit et distribution bottom-up."),
        ("Pourquoi pas juste utiliser ChatGPT ?", "ChatGPT est puissant mais hors flux. Assistia réduit le copier-coller, conserve le contexte visible et insère le brouillon directement dans l’outil de réponse."),
        ("La barrière à l’entrée est-elle suffisante ?", "Au départ, la barrière est faible. Elle se construit par données d’usage, templates métier, UX d’insertion robuste, marque francophone et distribution Chrome/SEO."),
        ("Pourquoi maintenant ?", "L’IA est acceptée par le marché, mais les utilisateurs veulent maintenant des assistants intégrés à des tâches concrètes plutôt que des chats génériques."),
        ("Est-ce assez grand pour du venture capital ?", "Oui si le produit dépasse le simple assistant Gmail et devient la couche de réponse professionnelle pour plusieurs canaux. Le marché large AI assistants est déjà multi-milliards."),
        ("Quel est le risque principal ?", "Distribution. Le produit doit prouver un canal d’acquisition rentable avant de lever lourd. C’est pourquoi le pre-seed doit financer une validation canal très disciplinée."),
        ("Pourquoi le prix est si bas ?", "Le prix d’entrée vise l’adoption prosumer et la conversion rapide. Le vrai levier ARR vient du Business/Team, du ton de marque et des templates partagés."),
        ("Qu’est-ce qui empêche les utilisateurs de churn après un mois ?", "L’usage doit devenir habituel : raccourcis, templates, historique léger, personnalisation du ton, et présence dans le champ de réponse."),
        ("Avez-vous besoin d’accéder aux emails complets ?", "Non au MVP. L’extension peut partir du texte visible/sélectionné. Cela réduit les risques OAuth, sécurité et review Google."),
        ("Quel serait un acquéreur naturel ?", "Grammarly/Superhuman, client email IA, suite CRM SMB, outil de productivité français/européen, ou acteur support/sales voulant ajouter une couche réponse conversationnelle.")
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        r = p.add_run(question)
        r.font.bold = True
        p2 = doc.add_paragraph(answer)
        p2.paragraph_format.left_indent = Cm(0.35)

    doc.add_heading("KPIs à tracker dès le lancement", level=2)
    add_bullets(doc, [
        "Installations Chrome, activation première réponse, réponses par WAU.",
        "Taux d’insertion du brouillon, régénérations, temps entre génération et insertion.",
        "Conversion free -> Pro, Pro -> Business, churn mensuel, NPS.",
        "CAC par canal, payback, LTV/CAC, marge brute par réponse générée.",
        "Top templates utilisés et cas d’usage convertissant le mieux."
    ])
    doc.add_heading("Glossaire", level=2)
    add_simple_table(doc, ["Terme", "Définition"], [
        ["MRR", "Monthly Recurring Revenue : revenu récurrent mensuel."],
        ["ARR", "Annual Recurring Revenue : MRR x 12."],
        ["CAC", "Customer Acquisition Cost : coût moyen pour acquérir un client payant."],
        ["LTV", "Lifetime Value : marge brute totale attendue sur un client."],
        ["Churn", "Pourcentage de clients qui annulent sur une période."],
        ["NRR", "Net Revenue Retention : revenu conservé + expansion - contraction/churn."],
        ["TAM", "Total Addressable Market : marché total théorique."],
        ["SAM", "Serviceable Addressable Market : partie du TAM atteignable avec la stratégie actuelle."],
        ["SOM", "Serviceable Obtainable Market : part réaliste capturable sur un horizon donné."],
        ["DSO", "Days Sales Outstanding : délai moyen d’encaissement client, surtout utile en B2B facturé."]
    ], [3.0, 13.7])

    doc.add_heading("Sections à compléter avant envoi", level=1)
    add_bullets(doc, [
        "Nom final : Assistia, Assistia Reply ou autre marque définitive.",
        "Bio fondateur : parcours, légitimité, avantage personnel.",
        "Traction réelle : waitlist, beta users, taux d’usage, retours clients.",
        "Montant exact recherché et instrument souhaité.",
        "Choix du marché initial : freelances, sales, agences, support client ou TPE généralistes.",
        "Données coûts réelles : OpenAI, Vercel, Supabase, Stripe, support."
    ])
    doc.add_heading("Sources à double-checker avant envoi", level=2)
    add_bullets(doc, [
        "Prix concurrents : ils changent vite, surtout Grammarly/Superhuman, Microsoft Copilot, HyperWrite et Compose AI.",
        "TAM : les rapports de marché privés utilisent des méthodologies différentes ; garder un wording prudent.",
        "SAM France : il repose sur une estimation de 20% des entreprises à forte intensité de communication client.",
        "Prévisions financières : elles sont des scénarios de travail, pas une promesse."
    ])
    doc.add_heading("3 améliorations recommandées du pitch", level=2)
    add_bullets(doc, [
        "Ajouter 5 captures produit : landing, panneau extension, génération, insertion dans Gmail, dashboard usage.",
        "Obtenir 10 citations beta users avec un avant/après de réponse client.",
        "Remplacer les estimations par des données propriétaires après 30 jours de beta : activation, usage, conversion, rétention."
    ])

    doc.add_heading("Sources", level=1)
    for name, title, url in SOURCES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name} - {title}: {url}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
