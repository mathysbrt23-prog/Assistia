from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "LECTURE.docx"
LOGO = ROOT / "public" / "assistia-logo.png"

SOURCES = [
    {
        "title": "Journal de bord",
        "path": ROOT / "JOURNAL_DE_BORD.md",
        "type": "markdown",
        "description": "Memoire commune du projet, decisions, historique et prochaines etapes.",
    },
    {
        "title": "Analyse produit - Assistia Reply",
        "path": ROOT / "docs" / "product" / "assistia-reply-analyse-produit.docx",
        "type": "docx",
        "description": "Analyse de faisabilite, MVP recommande, risques, pricing et roadmap.",
    },
    {
        "title": "Rapport concurrence - Assistia Reply",
        "path": ROOT / "docs" / "product" / "rapport-concurrence-assistia-reply.md",
        "type": "markdown",
        "description": "Veille concurrentielle sur les assistants de reponses integres aux emails et conversations.",
    },
    {
        "title": "README technique du projet",
        "path": ROOT / "README.md",
        "type": "markdown",
        "description": "Installation et configuration technique du projet Assistia.",
    },
    {
        "title": "Index des documents",
        "path": ROOT / "docs" / "README.md",
        "type": "markdown",
        "description": "Index court des documents conserves dans le dossier.",
    },
]


BLACK = "0B0D0F"
CHARCOAL = "171A1E"
GREEN = "25D366"
GREEN_DARK = "0B7A3F"
LIGHT = "F5F7F6"
BORDER = "D9E2DC"
TEXT = "171717"
MUTED = "5F6662"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=130, bottom=120, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Assistia - LECTURE | ")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Title", 30, "FFFFFF", 0, 8),
        ("Heading 1", 19, TEXT, 16, 7),
        ("Heading 2", 14, TEXT, 11, 5),
        ("Heading 3", 11.5, GREEN_DARK, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style.font.size = Pt(9.7)
        style.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, BLACK, "0")
    set_cell_margins(cell, top=560, start=520, bottom=560, end=520)
    if LOGO.exists():
        p = cell.paragraphs[0]
        p.add_run().add_picture(str(LOGO), width=Cm(2.0))
    else:
        p = cell.paragraphs[0]
        run = p.add_run("Assistia")
        run.font.color.rgb = rgb("FFFFFF")
        run.font.bold = True

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("LECTURE")
    run.font.name = "Arial"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = rgb("FFFFFF")

    p = cell.add_paragraph()
    run = p.add_run("Documents de reference du projet Assistia")
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = rgb("D7DEDA")

    p = cell.add_paragraph()
    run = p.add_run(
        "Journal de bord, analyse produit, rapport concurrence et notes techniques "
        "regroupes dans un seul fichier Word."
    )
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb("C2CBC6")

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("Derniere mise a jour : 28 avril 2026")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = rgb("98A29D")

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Contenu", level=1)
    for i, item in enumerate(SOURCES, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {item['title']}")
        r.font.bold = True
        p2 = doc.add_paragraph(item["description"])
        p2.paragraph_format.left_indent = Cm(0.55)
        p2.runs[0].font.color.rgb = rgb(MUTED)
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    return text.strip()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    inner = stripped.strip("|").replace(" ", "")
    return bool(inner) and all(ch in "-:|" for ch in inner)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        if not is_table_separator(line):
            cells = [clean_inline(cell) for cell in line.strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    table = doc.add_table(rows=1, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for row_index, row in enumerate(normalized):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for col_index, text in enumerate(row):
            cell = cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 90, 95, 90, 95)
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, CHARCOAL)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.color.rgb = rgb("FFFFFF")
                r.font.size = Pt(7.8 if max_cols >= 5 else 8.8)
            else:
                if col_index == 0:
                    set_cell_shading(cell, "F8FAF9")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(text)
                r.font.size = Pt(7.4 if max_cols >= 5 else 8.7)
    doc.add_paragraph()


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(0.25)
    right.width = Cm(15.5)
    set_cell_shading(left, GREEN)
    set_cell_shading(right, LIGHT)
    set_cell_border(left, GREEN, "0")
    set_cell_border(right, LIGHT, "0")
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 120, 160, 120, 160)
    p = right.paragraphs[0]
    r = p.add_run(clean_inline(text))
    r.font.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()


def add_markdown(doc: Document, path: Path, skip_first_h1: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    skipped_first_h1 = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Courier New"
            r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Cm(0.4)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if stripped.startswith("#"):
            level = min(3, max(1, len(stripped) - len(stripped.lstrip("#"))))
            text = clean_inline(stripped.lstrip("#").strip())
            if skip_first_h1 and level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                i += 1
                continue
            doc.add_heading(text, level=level)
            i += 1
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped.lstrip(">").strip())
            i += 1
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(bullet_match.group(1)))
            i += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean_inline(number_match.group(1)))
            i += 1
            continue

        paragraph_lines = [stripped]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith(">")
                or next_line.startswith("```")
                or re.match(r"^[-*]\s+", next_line)
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            j += 1
        doc.add_paragraph(clean_inline(" ".join(paragraph_lines)))
        i = j


def iter_docx_blocks(source: Document):
    body = source.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def add_docx_content(doc: Document, path: Path) -> None:
    source = Document(path)
    started = False
    para_by_el = {p._p: p for p in source.paragraphs}
    table_by_el = {t._tbl: t for t in source.tables}

    for kind, element in iter_docx_blocks(source):
        if kind == "paragraph":
            p = para_by_el.get(element)
            if p is None:
                continue
            text = p.text.strip()
            if not text:
                continue
            if text == "1. Verdict":
                started = True
            if not started:
                continue
            style_name = p.style.name if p.style else "Normal"
            if style_name.startswith("Heading 1"):
                doc.add_heading(text, level=1)
            elif style_name.startswith("Heading 2"):
                doc.add_heading(text, level=2)
            elif style_name.startswith("Heading 3"):
                doc.add_heading(text, level=3)
            elif style_name == "List Bullet":
                doc.add_paragraph(text, style="List Bullet")
            elif style_name == "List Number":
                doc.add_paragraph(text, style="List Number")
            else:
                doc.add_paragraph(text)
        elif kind == "table":
            if not started:
                continue
            table = table_by_el.get(element)
            if table is None:
                continue
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            add_table(doc, rows)


def add_landscape_section(doc: Document):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    set_footer(section)
    return section


def add_portrait_section(doc: Document):
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    set_footer(section)
    return section


def build() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    add_toc(doc)

    for index, item in enumerate(SOURCES, 1):
        path = item["path"]
        if not path.exists():
            continue
        if index > 1:
            doc.add_page_break()

        if item["type"] == "markdown":
            if item["title"].startswith("Rapport concurrence"):
                add_landscape_section(doc)
                doc.add_heading(f"{index}. {item['title']}", level=1)
                p = doc.add_paragraph(item["description"])
                p.runs[0].font.color.rgb = rgb(MUTED)
                add_markdown(doc, path, skip_first_h1=True)
                add_portrait_section(doc)
            else:
                doc.add_heading(f"{index}. {item['title']}", level=1)
                intro = doc.add_paragraph(item["description"])
                intro.runs[0].font.color.rgb = rgb(MUTED)
                add_markdown(doc, path, skip_first_h1=True)
        elif item["type"] == "docx":
            doc.add_heading(f"{index}. {item['title']}", level=1)
            intro = doc.add_paragraph(item["description"])
            intro.runs[0].font.color.rgb = rgb(MUTED)
            add_docx_content(doc, path)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
